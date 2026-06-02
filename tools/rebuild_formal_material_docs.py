from __future__ import annotations

import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\Pro\Pet2.0")
SRC_DOC = ROOT / "docs" / "大湾区资料" / "灵宠智芯-作品说明书（署名）.docx"
OUT_DIR = ROOT / "docs" / "机器人大赛"
ASSET_DIR = OUT_DIR / "_generated_assets"
RAW_DOC = OUT_DIR / "项目研究原始资料（图纸、图表、调查问卷等）.docx"
PHOTO_DOC = OUT_DIR / "项目研究活动照片.docx"

BLUE = RGBColor(31, 77, 120)
ACCENT = RGBColor(46, 116, 181)
INK = RGBColor(20, 30, 45)
MUTED = RGBColor(88, 96, 105)


def set_font(run, size: float | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell(cell, text: str, bold: bool = False, fill: str | None = None, center: bool = False) -> None:
    cell.text = ""
    if fill:
        shade(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_font(r, 9.2, bold, BLUE if bold else INK)


def apply_table_style(table) -> None:
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def setup_document(doc: Document, running_title: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.paragraph_format.space_after = Pt(0)
    hr = header.add_run(running_title)
    set_font(hr, 9, False, MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("第 ")
    set_font(fr, 9, False, MUTED)
    add_field(footer, "PAGE")
    fr2 = footer.add_run(" 页")
    set_font(fr2, 9, False, MUTED)


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, fld_end])


def cover(doc: Document, title: str, subtitle: str) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("第28届中国机器人及人工智能大赛")
    set_font(r, 14, True, BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(title)
    set_font(r2, 22, True, INK)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(subtitle)
    set_font(r3, 12, False, MUTED)
    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=2)
    apply_table_style(table)
    rows = [
        ("项目名称", "“灵宠智芯”——基于OpenHarmony的多模态全息数字虚拟宠物协同系统"),
        ("项目成员", "张梓悦、何荣庆、彭诗情"),
        ("指导教师", "何炜婷"),
        ("所在学校", "广东东软学院"),
        ("材料类型", title),
        ("整理日期", "2026年6月1日"),
    ]
    for row, (k, v) in zip(table.rows, rows):
        set_cell(row.cells[0], k, True, "F2F6FA", True)
        set_cell(row.cells[1], v)
    doc.add_page_break()


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    r = p.add_run(text)
    set_font(r, 15 if level == 1 else 12.5, True, BLUE if level == 1 else ACCENT)


def body(doc: Document, text: str, indent: bool = True) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    r = p.add_run(text)
    set_font(r, 10.5, False, INK)


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, 8.8, False, MUTED)


def picture(doc: Document, file_name: str, cap: str, width: float) -> None:
    path = ASSET_DIR / file_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    caption(doc, cap)


def picture_grid(doc: Document, items: list[tuple[str, str, str]], cols: int, image_width: float) -> None:
    table = doc.add_table(rows=math.ceil(len(items) / cols), cols=cols)
    apply_table_style(table)
    for row in table.rows:
        for cell in row.cells:
            shade(cell, "FFFFFF")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for idx, (file_name, title, note) in enumerate(items):
        cell = table.cell(idx // cols, idx % cols)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        path = ASSET_DIR / file_name
        if path.exists():
            p.add_run().add_picture(str(path), width=Cm(image_width))
        title_p = cell.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = title_p.add_run(title)
        set_font(tr, 8.8, True, BLUE)
        note_p = cell.add_paragraph()
        nr = note_p.add_run(note)
        set_font(nr, 8.2, False, MUTED)
    doc.add_paragraph()


def src_table(index: int) -> list[list[str]]:
    table = Document(SRC_DOC).tables[index]
    return [[cell.text.strip().replace("\n", " / ") for cell in row.cells] for row in table.rows]


def data_table(doc: Document, title: str, rows: list[list[str]]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, 10.2, True, BLUE)
    table = doc.add_table(rows=1, cols=len(rows[0]))
    apply_table_style(table)
    for cell, text in zip(table.rows[0].cells, rows[0]):
        set_cell(cell, text, True, "F2F6FA", True)
    for row_data in rows[1:]:
        row = table.add_row().cells
        for cell, text in zip(row, row_data):
            set_cell(cell, text)
    doc.add_paragraph()


def index_table(doc: Document, title: str, rows: list[tuple[str, str, str, str]]) -> None:
    heading(doc, title, 1)
    table = doc.add_table(rows=1, cols=4)
    apply_table_style(table)
    for cell, text in zip(table.rows[0].cells, ["编号", "资料名称", "内容说明", "证明用途"]):
        set_cell(cell, text, True, "F2F6FA", True)
    for row_data in rows:
        row = table.add_row().cells
        for cell, text in zip(row, row_data):
            set_cell(cell, text, center=len(text) <= 8)
    doc.add_paragraph()


def build_raw() -> None:
    doc = Document()
    setup_document(doc, "项目研究原始资料")
    cover(doc, "项目研究原始资料", "图纸、图表、调查问卷等")
    index_table(doc, "一、资料编制说明与索引", [
        ("R-01", "系统图纸", "系统框架图、三端协同数据流图、硬件平台图、芯片逻辑图", "说明总体技术路线与硬件关系"),
        ("R-02", "技术表格", "开发环境与技术选型、硬件设计、传感器反馈、功能性能测试", "支撑技术实现说明"),
        ("R-03", "界面截图", "主页面、宠物选择、档案、AI对话、喂食、传感器页面", "证明软件功能完成情况"),
        ("R-04", "调查问卷", "用户需求调查题项、选项设计与结果归纳图", "支撑需求分析"),
        ("R-05", "研究计划", "前期、中期、后期计划表与研究时间轴", "证明研究过程完整性"),
    ])
    body(doc, "本文件以作品说明书为依据，对项目研究过程中形成的图纸、图表、界面截图、调查问卷和阶段计划进行归档。材料采用统一编号，图表标题置于图表附近，说明文字保持客观、可追溯。")

    heading(doc, "二、系统图纸与技术路线", 1)
    picture(doc, "image3.png", "图R-1 系统三端协同架构图", 14.0)
    picture(doc, "三端协同数据流示意图.png", "图R-2 三端协同数据流示意图", 14.2)
    data_table(doc, "表R-1 开发环境与技术选型表", src_table(3))
    data_table(doc, "表R-2 硬件设计表", src_table(2))

    heading(doc, "三、问题分析、功能设计与测试结果", 1)
    data_table(doc, "表R-3 现有产品缺陷表", src_table(1))
    picture(doc, "能力矩阵图.png", "图R-3 核心能力矩阵", 14.2)
    data_table(doc, "表R-4 传感器功能与反馈机制表", src_table(5))
    data_table(doc, "表R-5 功能与性能表现表", src_table(6))

    heading(doc, "四、软件界面原始截图", 1)
    picture_grid(doc, [
        ("image10.png", "图R-4 程序主页面", "展示宠物状态、交互入口与主控布局。"),
        ("image5.png", "图R-5 宠物选择与命名界面", "记录宠物初始化和命名流程。"),
        ("image7.png", "图R-6 宠物档案界面", "展示成长记录、互动记忆和属性信息。"),
        ("image8.png", "图R-7 AI对话界面", "展示自然语言互动记录。"),
        ("image9.png", "图R-8 喂食功能界面", "展示用户行为与状态变化反馈。"),
        ("image15.png", "图R-9 温湿度传感器界面", "展示南向感知数据可视化。"),
    ], 3, 4.7)

    heading(doc, "五、硬件联调与全息展示资料", 1)
    picture_grid(doc, [
        ("image18.png", "图R-10 开发板硬件平台图", "记录开发板和外设接口关系。"),
        ("image17.png", "图R-11 WS63V100芯片逻辑框图", "说明核心模组引脚与硬件基础。"),
        ("image1.jpeg", "图R-12 实物联动环境", "展示手机端、开发板与全息装置联动。"),
        ("image19.png", "图R-13 全息显示测试照片一", "记录全息宠物显示效果。"),
        ("image20.png", "图R-14 全息显示测试照片二", "记录不同观察角度下的显示状态。"),
        ("image21.png", "图R-15 全息显示测试照片三", "记录宠物姿态和展示稳定性。"),
    ], 2, 6.8)

    heading(doc, "六、用户需求调查问卷样表", 1)
    body(doc, "问卷围绕虚拟宠物陪伴、长期记忆、全息展示、跨端协同和隐私运行能力设计。正式发放时，可在本样表基础上补充样本数量、回收时间、有效问卷数和统计口径。")
    picture(doc, "用户需求调研归纳图.png", "图R-16 用户需求调研归纳图", 14.2)
    table = doc.add_table(rows=1, cols=4)
    apply_table_style(table)
    for cell, text in zip(table.rows[0].cells, ["题号", "调查问题", "选项设计", "记录口径"]):
        set_cell(cell, text, True, "F2F6FA", True)
    rows = [
        ("Q1", "是否使用过电子宠物、虚拟角色或AI陪伴产品？", "是 / 否 / 偶尔接触", "单选"),
        ("Q2", "最看重虚拟宠物的哪些能力？", "情绪反馈 / 长期记忆 / 成长档案 / 全息展示 / 环境提醒", "多选"),
        ("Q3", "是否愿意通过手机端与桌面全息设备共同使用虚拟宠物？", "非常愿意 / 比较愿意 / 一般 / 不愿意", "单选"),
        ("Q4", "认为现有虚拟宠物最需要解决的痛点是什么？", "交互单一 / 无记忆 / 缺少真实陪伴 / 设备割裂 / 其他", "多选"),
        ("Q5", "对宠物个性成长、梦境生成和行为反馈的接受度如何？", "非常期待 / 比较期待 / 一般 / 暂不需要", "单选"),
        ("Q6", "是否关注本地运行、隐私保护和低硬件依赖能力？", "非常关注 / 比较关注 / 一般 / 不关注", "单选"),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for cell, text in zip(row, row_data):
            set_cell(cell, text, center=len(text) <= 4)

    heading(doc, "七、研究计划与阶段成果", 1)
    picture(doc, "研究活动时间轴.png", "图R-17 项目研究活动时间轴", 14.2)
    data_table(doc, "表R-6 前期计划表", src_table(7))
    data_table(doc, "表R-7 中期计划表", src_table(8))
    data_table(doc, "表R-8 后期计划表", src_table(9))
    doc.save(RAW_DOC)


def build_photo() -> None:
    doc = Document()
    setup_document(doc, "项目研究活动照片")
    cover(doc, "项目研究活动照片", "开发、联调、测试与展示记录")
    index_table(doc, "一、活动照片索引", [
        ("P-01", "实物联调环境", "手机端、开发板与全息装置同步展示", "三端联动证明"),
        ("P-02", "开发板硬件平台", "开发板、传感器与外设接口记录", "硬件联调证明"),
        ("P-03", "全息显示测试", "不同角度和光照下的全息宠物显示", "展示效果证明"),
        ("P-04", "移动端功能验证", "主页面、选择、档案、对话、喂食与传感器页面", "软件功能证明"),
        ("P-05", "测试与材料整理", "性能指标、活动记录和附件材料整理", "过程完整性证明"),
    ])
    body(doc, "本文件用于记录项目研究活动中的关键照片和截图。每组照片均配套说明活动阶段、记录内容和对应成果，便于评审人员理解项目从设计、开发、联调到测试整理的完整过程。")

    heading(doc, "二、硬件联调与全息显示测试", 1)
    picture_grid(doc, [
        ("image1.jpeg", "图P-1 实物联调环境", "记录手机端、开发板和全息显示端同时运行的场景。"),
        ("image18.png", "图P-2 开发板硬件平台", "记录开发板接口、传感器接入和硬件基础。"),
        ("image19.png", "图P-3 全息显示测试一", "验证宠物形象在全息装置中的可见性。"),
        ("image20.png", "图P-4 全息显示测试二", "记录侧向观察下的显示效果。"),
        ("image21.png", "图P-5 全息显示测试三", "记录宠物动作姿态和展示稳定性。"),
    ], 2, 6.8)

    heading(doc, "三、移动端功能验证记录", 1)
    picture_grid(doc, [
        ("image10.png", "图P-6 程序主页面", "验证主控入口、宠物状态和交互面板。"),
        ("image5.png", "图P-7 宠物选择与命名", "验证初始化流程和用户输入保存。"),
        ("image7.png", "图P-8 宠物档案界面", "验证成长档案和互动记忆展示。"),
        ("image8.png", "图P-9 AI对话界面", "验证连续对话和交互记录。"),
        ("image9.png", "图P-10 喂食功能界面", "验证行为触发与状态反馈。"),
        ("image15.png", "图P-11 传感器监测界面", "验证温湿度等环境数据展示。"),
    ], 3, 4.7)

    heading(doc, "四、研究活动过程记录", 1)
    table = doc.add_table(rows=1, cols=5)
    apply_table_style(table)
    for cell, text in zip(table.rows[0].cells, ["阶段", "时间", "活动内容", "记录要点", "成果材料"]):
        set_cell(cell, text, True, "F2F6FA", True)
    rows = [
        ("前期调研", "2026年5月上旬", "需求分析、竞品比较、技术路线确认", "明确情绪陪伴、长期记忆、跨端协同和全息展示价值", "需求归纳图、产品缺陷表"),
        ("原型开发", "2026年5月中旬", "ArkTS/ArkUI页面与交互流程实现", "验证主页面、喂食、档案、AI对话、传感器页面", "移动端功能截图"),
        ("硬件联调", "2026年5月中下旬", "开发板、传感器与全息显示装置连接", "验证感知数据回传、设备同步和展示链路", "硬件平台图、实物照片"),
        ("测试完善", "2026年5月下旬", "功能、性能、兼容性和展示效果检查", "记录帧率、延迟、启动时间和本地存储表现", "测试表、全息测试照片"),
        ("材料整理", "2026年6月1日", "申报材料、测试附件、研究资料和活动照片归档", "统一编号、补充图表、核对附件清单", "最终Word材料包"),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for cell, text in zip(row, row_data):
            set_cell(cell, text, center=len(text) <= 8)

    heading(doc, "五、照片真实性与用途说明", 1)
    body(doc, "本文件中的照片与截图均围绕作品说明书所述功能、硬件和展示环节整理，用于说明项目研究、开发、联调、测试和展示过程。材料提交时可与项目研究报告、系统测试报告和多场景测试验证数据配套使用。")
    doc.save(PHOTO_DOC)


def main() -> None:
    build_raw()
    build_photo()
    print(RAW_DOC)
    print(PHOTO_DOC)


if __name__ == "__main__":
    main()
