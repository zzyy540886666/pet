from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


DOC_PATH = Path(r"D:\Pro\Pet2.0\docs\机器人大赛\附件：人工智能创新赛报名表、查新报告.docx")


def format_cell_text(cell, lines: list[str]) -> None:
    cell.text = ""
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(line)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(10)
        if index == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(31, 77, 120)


def main() -> None:
    document = Document(DOC_PATH)
    table = document.tables[5]

    attachment_lines = [
        "八．附件清单",
        "1．密切相关文献附件1：《一种控制虚拟宠物的方法及智能投影设备》，出处：国家知识产权局专利文献；附件形式：原文复制件。",
        "2．密切相关文献附件2：《Mixed Reality-Based Interaction between Human and Virtual Cat for Mental Stress Management》，出处：Sensors，2022；附件形式：文献原文或摘要页。",
        "3．一般相关文献附件3：《基于NB-IoT网络的智能宠物项圈》，出处：国家知识产权局专利文献；附件形式：原文复制件或摘要页。",
        "4．一般相关文献附件4：《OpenHarmony轻智能设备迎来“电子宠物机”应用，让陪伴触手可及！》，出处：CSDN 技术博客；附件形式：网页题录及内容摘要。",
        "5．补充说明：系统架构图、核心能力矩阵、三端协同数据流图、用户需求调研归纳图、硬件平台图、全息显示测试图和研究活动时间轴已插入查新报告正文，不作为单独附件重复列出。",
    ]
    format_cell_text(table.rows[8].cells[0], attachment_lines)

    remark_lines = [
        "备注",
        "本附件清单与“五．检索结果”中提供附件（4）份一致；项目研究报告、项目研究原始资料、项目研究活动照片等申报材料已在“C、项目申报材料统计”中单独列明。",
    ]
    format_cell_text(table.rows[9].cells[0], remark_lines)

    document.save(DOC_PATH)
    print(DOC_PATH)


if __name__ == "__main__":
    main()
