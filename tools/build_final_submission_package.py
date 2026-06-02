from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(r"D:\Pro\Pet2.0")
SRC_DIR = ROOT / "docs" / "机器人大赛"
DEST_DIR = SRC_DIR / "提交材料_基于OpenHarmony与人工智能记忆增强的多模态全息数字虚拟宠物"

PROJECT_NAME = "基于OpenHarmony与人工智能记忆增强的多模态全息数字虚拟宠物"

SOURCE_FILES = {
    "01_人工智能创新赛报名表_查新报告.docx": SRC_DIR / "附件：人工智能创新赛报名表、查新报告.docx",
    "02_项目研究报告.docx": SRC_DIR / "项目研究报告.docx",
    "03_项目研究原始资料.docx": SRC_DIR / "项目研究原始资料（图纸、图表、调查问卷等）.docx",
    "04_项目研究活动照片.docx": SRC_DIR / "项目研究活动照片.docx",
    "05_多场景测试验证数据.docx": SRC_DIR / "多场景测试验证数据（实验室测试、家庭测试）.docx",
    "06_系统测试报告.docx": SRC_DIR / "系统测试报告（功能测试与性能测试结果）.docx",
}


TITLE_REPLACEMENTS = [
    ("“灵宠智芯”——基于OpenHarmony的多模态全息数字虚拟宠物协同系统", f"“{PROJECT_NAME}”"),
    ("“灵宠智芯”--基于OpenHarmony的多模态全息数字虚拟宠物协同系统", f"“{PROJECT_NAME}”"),
    ("“灵宠智芯”——基于OpenHarmony的多模态全息数字虚拟宠物协同系统", f"“{PROJECT_NAME}”"),
    ("灵宠智芯项目附件", f"{PROJECT_NAME}项目附件"),
    ("灵宠智芯", PROJECT_NAME),
]


def set_font(run, size: float | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def replace_text_in_paragraph(paragraph, replacements: list[tuple[str, str]]) -> None:
    for run in paragraph.runs:
        text = run.text
        for old, new in replacements:
            if old in text:
                text = text.replace(old, new)
        run.text = text


def replace_text_in_table(table, replacements: list[tuple[str, str]]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)


def replace_text_everywhere(doc: Document, replacements: list[tuple[str, str]]) -> None:
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        replace_text_in_table(table, replacements)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        for table in section.header.tables:
            replace_text_in_table(table, replacements)
        for paragraph in section.footer.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)
        for table in section.footer.tables:
            replace_text_in_table(table, replacements)


def find_first_paragraph_index(doc: Document, needle: str) -> int | None:
    for idx, paragraph in enumerate(doc.paragraphs):
        if needle in paragraph.text:
            return idx
    return None


def insert_paragraph_after(paragraph, text: str, size: float, bold: bool, color: RGBColor, align=WD_ALIGN_PARAGRAPH.CENTER):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    new_paragraph = Paragraph(new_p, paragraph._parent)
    new_paragraph.alignment = align
    new_paragraph.paragraph_format.space_after = Pt(0)
    run = new_paragraph.add_run(text)
    set_font(run, size, bold, color)
    return new_paragraph


def insert_project_name_cover(doc: Document, anchor_text: str, after_text: str | None = None) -> None:
    idx = find_first_paragraph_index(doc, anchor_text)
    if idx is None:
        return
    if after_text is None:
        insert_after = doc.paragraphs[idx]
    else:
        target_idx = find_first_paragraph_index(doc, after_text)
        if target_idx is None:
            insert_after = doc.paragraphs[idx]
        else:
            insert_after = doc.paragraphs[target_idx]
    insert_paragraph_after(insert_after, PROJECT_NAME, 18, True, RGBColor(31, 77, 120))


def set_page_defaults(doc: Document) -> None:
    for section in doc.sections:
        section.header_distance = Pt(28)
        section.footer_distance = Pt(28)
        for paragraph in section.header.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, 9, False, RGBColor(88, 96, 105))
        for paragraph in section.footer.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_font(run, 9, False, RGBColor(88, 96, 105))


def update_document(src: Path, dest: Path, mode: str) -> None:
    doc = Document(src)
    replace_text_everywhere(doc, TITLE_REPLACEMENTS)

    if mode == "report":
        idx = find_first_paragraph_index(doc, "参赛作品说明书")
        if idx is not None:
            doc.paragraphs[idx].text = "项目研究报告"
            for run in doc.paragraphs[idx].runs:
                set_font(run, 16, True, RGBColor(31, 77, 120))
            doc.paragraphs[idx].alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_idx = find_first_paragraph_index(doc, PROJECT_NAME)
        if title_idx is not None:
            doc.paragraphs[title_idx].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif mode == "raw":
        insert_project_name_cover(doc, "第28届中国机器人及人工智能大赛", after_text="第28届中国机器人及人工智能大赛")
    elif mode == "photo":
        insert_project_name_cover(doc, "第28届中国机器人及人工智能大赛", after_text="第28届中国机器人及人工智能大赛")
    elif mode == "registration":
        title_idx = find_first_paragraph_index(doc, "人工智能创新赛报名表")
        if title_idx is not None:
            insert_paragraph_after(doc.paragraphs[title_idx], PROJECT_NAME, 14, True, RGBColor(31, 77, 120))
        if len(doc.tables) > 2:
            doc.tables[2].cell(0, 0).text = (
                "1．人工智能创新比赛项目报名表1份\n"
                "2．项目查新报告1份（图示、关键技术表格已纳入查新报告正文）\n"
                "3．项目研究报告1份\n"
                "4．项目研究原始资料（图纸、图表、调查问卷等）1份\n"
                "5．项目研究活动照片1份\n"
                "6．多场景测试验证数据1份\n"
                "7．系统测试报告1份"
            )
        if len(doc.tables) > 5:
            doc.tables[5].rows[8].cells[0].text = (
                "八．附件清单\n"
                "1．多场景测试验证数据（实验室测试、家庭测试）.docx\n"
                "2．系统测试报告（功能测试与性能测试结果）.docx\n"
                "3．项目研究报告.docx\n"
                "4．项目研究原始资料（图纸、图表、调查问卷等）.docx\n"
                "5．项目研究活动照片.docx"
            )
            doc.tables[5].rows[9].cells[0].text = (
                "备注\n"
                "上述附件均为最终提交包中的真实 Word 文件。"
            )
    elif mode in {"scene", "system"}:
        # The subtitle line already carries the project name after replacement.
        pass

    set_page_defaults(doc)
    doc.save(dest)


def main() -> None:
    DEST_DIR.mkdir(parents=True, exist_ok=True)
    for dest_name, src_path in SOURCE_FILES.items():
        if "项目研究报告" in dest_name:
            mode = "report"
        elif "原始资料" in dest_name:
            mode = "raw"
        elif "活动照片" in dest_name:
            mode = "photo"
        elif "报名表" in dest_name:
            mode = "registration"
        elif "多场景测试" in dest_name:
            mode = "scene"
        else:
            mode = "system"
        update_document(src_path, DEST_DIR / dest_name, mode)

    manifest = DEST_DIR / "00_材料清单.txt"
    manifest.write_text(
        "\n".join(
            [
                "提交材料目录",
                f"1. 01_人工智能创新赛报名表_查新报告.docx",
                f"2. 02_项目研究报告.docx",
                f"3. 03_项目研究原始资料.docx",
                f"4. 04_项目研究活动照片.docx",
                f"5. 05_多场景测试验证数据.docx",
                f"6. 06_系统测试报告.docx",
            ]
        ),
        encoding="utf-8",
    )
    print(DEST_DIR)


if __name__ == "__main__":
    main()
