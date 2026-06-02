from __future__ import annotations

import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\Pro\Pet2.0")
OUT_DIR = ROOT / "docs" / "机器人大赛"
ASSET_DIR = OUT_DIR / "_generated_assets"
REG_DOC = OUT_DIR / "附件：人工智能创新赛报名表、查新报告.docx"
RAW_DOC = OUT_DIR / "项目研究原始资料（图纸、图表、调查问卷等）.docx"
PHOTO_DOC = OUT_DIR / "项目研究活动照片.docx"

ACCENT = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(85, 85, 85)


def font_path() -> str:
    for path in [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]:
        if Path(path).exists():
            return path
    return r"C:\Windows\Fonts\arial.ttf"


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    if bold and Path(r"C:\Windows\Fonts\msyhbd.ttc").exists():
        return ImageFont.truetype(r"C:\Windows\Fonts\msyhbd.ttc", size)
    return ImageFont.truetype(font_path(), size)


def set_font(run, size: float | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    if fill:
        shade(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 12 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_font(r, 9.5, bold, ACCENT if bold else None)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    r = p.add_run(text)
    set_font(r, 16 if level == 1 else 13, True, BLUE if level <= 2 else ACCENT)


def add_body(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, 10.5, bold, None)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run(text)
    set_font(r, 9, False, MUTED)


def add_picture(doc: Document, file_name: str, caption: str, width_cm: float) -> None:
    path = ASSET_DIR / file_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_picture_to_cell(cell, file_name: str, caption: str, width_cm: float) -> None:
    path = ASSET_DIR / file_name
    if not path.exists():
        return
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = cell.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_font(r, 8.5, False, MUTED)


def add_picture_grid(doc: Document, items: list[tuple[str, str]], cols: int, width_cm: float) -> None:
    rows = math.ceil(len(items) / cols)
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            shade(cell, "FFFFFF")
    for i, (file_name, title) in enumerate(items):
        cell = table.cell(i // cols, i % cols)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        path = ASSET_DIR / file_name
        if path.exists():
            p.add_run().add_picture(str(path), width=Cm(width_cm - 0.35))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(title)
        set_font(r, 8.8, True, INK)
    doc.add_paragraph()


def remove_tail_supplement(doc: Document) -> None:
    body = doc.element.body
    children = list(body)
    start = None
    for idx, child in enumerate(children):
        if child.tag == qn("w:p") and "附件：项目图示补充材料" in "".join(child.itertext()):
            start = idx
            break
    if start is None:
        return
    for child in children[start:]:
        body.remove(child)


def update_material_stats(doc: Document) -> None:
    text = (
        "1．人工智能创新比赛项目报名表1份\n"
        "2．项目研究报告1份，项目查新报告1份（图示、系统框架、能力矩阵已纳入查新报告正文）\n"
        "3．项目研究原始资料（图纸、图表、调查问卷等）1份\n"
        "4．项目研究活动照片1份\n"
        "5．附件材料：项目说明书、演示视频、承诺书、实物与全息展示图片等按竞赛要求另行提交。"
    )
    if len(doc.tables) > 2:
        doc.tables[2].cell(0, 0).text = text
    if len(doc.tables) > 5 and len(doc.tables[5].rows) > 8:
        doc.tables[5].rows[8].cells[1].text = (
            "图示材料已插入本查新报告正文；其他附件包括项目研究报告、项目研究原始资料、项目研究活动照片及相关说明文件。"
        )


def insert_figures_into_report(doc: Document) -> None:
    if len(doc.tables) <= 5:
        return
    table = doc.tables[5]
    innovation_cell = table.rows[2].cells[0]
    point_cell = table.rows[3].cells[0]
    strategy_cell = table.rows[4].cells[0]
    result_cell = table.rows[5].cells[0]
    conclusion_cell = table.rows[6].cells[0]

    add_picture_to_cell(innovation_cell, "image3.png", "图1 系统三端协同架构图", 11.5)
    add_picture_to_cell(innovation_cell, "能力矩阵图.png", "图2 项目核心能力矩阵", 11.5)
    add_picture_to_cell(point_cell, "三端协同数据流示意图.png", "图3 三端协同数据流示意图", 11.5)
    add_picture_to_cell(strategy_cell, "用户需求调研归纳图.png", "图4 用户需求调研归纳图", 11.5)
    add_picture_to_cell(result_cell, "image18.png", "图5 开发板硬件平台图", 9.6)
    add_picture_to_cell(result_cell, "image19.png", "图6 全息显示测试照片", 9.6)
    add_picture_to_cell(conclusion_cell, "研究活动时间轴.png", "图7 项目研究活动时间轴", 11.5)


def polish_registration_doc() -> None:
    doc = Document(REG_DOC)
    remove_tail_supplement(doc)
    update_material_stats(doc)
    already_inserted = any("图1 系统三端协同架构图" in p.text for p in doc.paragraphs)
    if not already_inserted:
        insert_figures_into_report(doc)
    doc.save(REG_DOC)


def create_cover_banner(path: Path, title: str, subtitle: str) -> None:
    w, h = 1600, 520
    im = Image.new("RGB", (w, h), "#F7FAFD")
    draw = ImageDraw.Draw(im)
    title_font = pil_font(58, True)
    sub_font = pil_font(30)
    small_font = pil_font(24)
    draw.rounded_rectangle((60, 55, 1540, 465), radius=36, fill="#E8EEF5", outline="#7DA6D3", width=4)
    draw.rectangle((60, 55, 205, 465), fill="#2E74B5")
    draw.text((260, 150), title, font=title_font, fill="#0B2545")
    draw.text((265, 235), subtitle, font=sub_font, fill="#333333")
    draw.text((265, 330), "灵宠智芯 | OpenHarmony 多模态全息数字虚拟宠物协同系统", font=small_font, fill="#555555")
    im.save(path)


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)


def add_cover(doc: Document, title: str, subtitle: str, banner_name: str) -> None:
    banner = ASSET_DIR / banner_name
    create_cover_banner(banner, title, subtitle)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(banner), width=Cm(16.2))
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows = [
        ("项目名称", "“灵宠智芯”——基于OpenHarmony的多模态全息数字虚拟宠物协同系统"),
        ("参赛方向", "人工智能创新赛"),
        ("资料来源", "《灵宠智芯-作品说明书（署名）》及补充整理图表"),
        ("整理日期", "2026年6月1日"),
    ]
    for row, (left, right) in zip(table.rows, rows):
        cell_text(row.cells[0], left, True, "F2F4F7")
        cell_text(row.cells[1], right)
    doc.add_page_break()


def add_overview_cards(doc: Document, cards: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    for idx, (title, body) in enumerate(cards):
        cell = table.cell(idx // 2, idx % 2)
        shade(cell, "E8EEF5" if idx % 2 == 0 else "F2F7FC")
        p = cell.paragraphs[0]
        r = p.add_run(title)
        set_font(r, 11.5, True, ACCENT)
        p2 = cell.add_paragraph()
        r2 = p2.add_run(body)
        set_font(r2, 9.5, False, None)
    doc.add_paragraph()


def build_raw_doc() -> None:
    doc = Document()
    setup_doc(doc)
    add_cover(doc, "项目研究原始资料", "图纸、图表、调查问卷等归档材料", "原始资料封面.png")
    add_heading(doc, "一、资料概览", 1)
    add_overview_cards(doc, [
        ("图纸资料", "系统框架图、硬件平台图、芯片逻辑图、三端协同数据流图。"),
        ("图表资料", "核心能力矩阵、研究活动时间轴、用户需求调研归纳图。"),
        ("界面截图", "主页面、宠物选择、宠物档案、AI 对话、喂食和传感器界面。"),
        ("问卷样表", "围绕陪伴需求、长期记忆、跨端协同、全息展示和隐私运行设计。"),
    ])
    add_body(doc, "本文件依据作品说明书中的图片、技术路线、系统实现和测试内容整理，用于支撑项目申报材料中的研究过程证明、图纸图表归档和需求调研说明。")

    add_heading(doc, "二、技术图纸与图表", 1)
    add_picture(doc, "image3.png", "图1 系统三端协同架构图", 14.0)
    add_picture(doc, "三端协同数据流示意图.png", "图2 三端协同数据流示意图", 14.2)
    add_picture(doc, "能力矩阵图.png", "图3 核心能力矩阵", 14.2)
    add_picture(doc, "研究活动时间轴.png", "图4 项目研究活动时间轴", 14.2)

    add_heading(doc, "三、软件界面原始截图", 1)
    add_picture_grid(doc, [
        ("image10.png", "程序主页面"),
        ("image5.png", "宠物选择与命名"),
        ("image7.png", "宠物档案"),
        ("image8.png", "AI 对话"),
        ("image9.png", "喂食与状态反馈"),
        ("image15.png", "温湿度传感器"),
    ], cols=3, width_cm=5.0)

    add_heading(doc, "四、硬件与全息展示资料", 1)
    add_picture_grid(doc, [
        ("image18.png", "开发板硬件平台"),
        ("image17.png", "芯片逻辑框图"),
        ("image1.jpeg", "实物与全息联动"),
        ("image19.png", "全息显示测试一"),
        ("image20.png", "全息显示测试二"),
        ("image21.png", "全息显示测试三"),
    ], cols=2, width_cm=7.3)

    add_heading(doc, "五、用户需求调查问卷样表", 1)
    add_picture(doc, "用户需求调研归纳图.png", "图5 用户需求调研归纳图", 14.2)
    questions = [
        "你是否使用过电子宠物、虚拟角色或 AI 陪伴产品？",
        "你最看重虚拟宠物的哪些能力：情绪反馈、长期记忆、成长档案、全息展示、环境提醒？",
        "你是否愿意通过手机端与桌面全息设备共同使用虚拟宠物？",
        "你认为虚拟宠物最需要解决的痛点是什么？",
        "你对宠物个性成长、梦境生成和行为反馈的接受度如何？",
        "你是否关注本地运行、隐私保护和低硬件依赖能力？",
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["题号", "问题", "选项/记录"]):
        cell_text(cell, text, True, "F2F4F7")
    for idx, question in enumerate(questions, 1):
        cells = table.add_row().cells
        cell_text(cells[0], str(idx))
        cell_text(cells[1], question)
        cell_text(cells[2], "非常需要 / 比较需要 / 一般 / 暂不需要")
    doc.save(RAW_DOC)


def build_photo_doc() -> None:
    doc = Document()
    setup_doc(doc)
    add_cover(doc, "项目研究活动照片", "开发、联调、测试与展示记录", "活动照片封面.png")
    add_heading(doc, "一、活动记录概览", 1)
    add_overview_cards(doc, [
        ("需求与方案", "围绕用户陪伴痛点、产品定位和技术路线完成前期规划。"),
        ("软件开发", "完成宠物主界面、AI 对话、喂食互动、档案和传感器页面。"),
        ("硬件联调", "完成开发板、传感器和全息显示装置的联动验证。"),
        ("展示测试", "记录全息宠物显示效果、姿态表现和运行稳定性。"),
    ])
    add_heading(doc, "二、硬件联调与全息显示测试", 1)
    add_picture_grid(doc, [
        ("image1.jpeg", "实物联调环境"),
        ("image18.png", "开发板硬件平台"),
        ("image19.png", "全息显示测试一"),
        ("image20.png", "全息显示测试二"),
        ("image21.png", "全息显示测试三"),
    ], cols=2, width_cm=7.3)
    add_heading(doc, "三、移动端功能验证记录", 1)
    add_picture_grid(doc, [
        ("image10.png", "程序主页面"),
        ("image5.png", "宠物选择与命名"),
        ("image7.png", "宠物档案界面"),
        ("image8.png", "AI 对话界面"),
        ("image9.png", "喂食功能界面"),
        ("image15.png", "传感器监测界面"),
    ], cols=3, width_cm=5.0)
    add_heading(doc, "四、研究活动记录表", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["阶段", "活动内容", "记录要点", "成果"]):
        cell_text(cell, text, True, "F2F4F7")
    rows = [
        ("前期调研", "需求分析、竞品比较、技术路线确认", "明确情绪陪伴、长期记忆和跨设备协同价值", "项目总体方案"),
        ("原型开发", "ArkTS/ArkUI 页面与交互流程实现", "验证主页面、喂食、档案和 AI 对话闭环", "软件功能原型"),
        ("硬件联调", "开发板、传感器与显示装置连接", "验证南向感知数据回传和设备协同", "硬件接口资料"),
        ("展示测试", "全息宠物显示效果记录", "验证浮空视觉、姿态显示和演示稳定性", "活动照片与演示素材"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            cell_text(cell, text)
    doc.save(PHOTO_DOC)


def main() -> None:
    polish_registration_doc()
    build_raw_doc()
    build_photo_doc()
    print("updated", REG_DOC)
    print("polished", RAW_DOC)
    print("polished", PHOTO_DOC)


if __name__ == "__main__":
    main()
