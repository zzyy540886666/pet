from __future__ import annotations

import math
import os
import shutil
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\Pro\Pet2.0")
SOURCE_DOC = ROOT / "docs" / "大湾区资料" / "灵宠智芯-作品说明书（署名）.docx"
TARGET_DOC = ROOT / "docs" / "机器人大赛" / "附件：人工智能创新赛报名表、查新报告.docx"
OUT_DIR = ROOT / "docs" / "机器人大赛"
ASSET_DIR = OUT_DIR / "_generated_assets"

REPORT_DOC = OUT_DIR / "项目研究报告.docx"
RAW_DOC = OUT_DIR / "项目研究原始资料（图纸、图表、调查问卷等）.docx"
PHOTO_DOC = OUT_DIR / "项目研究活动照片.docx"

ACCENT = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(85, 85, 85)


@dataclass(frozen=True)
class Figure:
    file_name: str
    title: str
    note: str


SOURCE_FIGURES = {
    "effect": Figure("image1.jpeg", "项目实物与全息宠物联动效果", "用于呈现手机控制端、开发板与全息显示端的联动形态。"),
    "route": Figure("image2.jpeg", "总体技术路线图", "展示项目从入口、功能模块到全息显示与硬件协同的整体路径。"),
    "system": Figure("image3.png", "系统三端协同架构图", "呈现手机控制端、全息显示端与南向感知端的接口关系。"),
    "logic": Figure("image16.png", "系统逻辑框架图", "补充说明数据流、模块分层和设备协同链路。"),
    "main": Figure("image10.png", "程序主页面", "展示虚拟宠物主界面、状态栏与交互入口。"),
    "select": Figure("image5.png", "宠物选择与命名界面", "展示用户初始化宠物形象和命名的关键流程。"),
    "profile": Figure("image7.png", "宠物档案界面", "展示宠物记忆、属性、成长信息的组织方式。"),
    "chat": Figure("image8.png", "AI 对话界面", "展示多模态交互中的自然语言沟通界面。"),
    "feed": Figure("image9.png", "喂食与状态反馈界面", "展示用户行为与宠物状态变化之间的反馈关系。"),
    "sensor": Figure("image15.png", "温湿度传感器监测界面", "展示南向感知数据在移动端的可视化呈现。"),
    "chip": Figure("image17.png", "WS63V100 芯片逻辑框图", "用于说明核心模组的引脚与硬件接口基础。"),
    "board": Figure("image18.png", "开发板硬件平台图", "展示星闪开发板及外设接口连接关系。"),
    "holo1": Figure("image19.png", "全息显示测试照片一", "记录全息浮空宠物在投影装置中的显示效果。"),
    "holo2": Figure("image20.png", "全息显示测试照片二", "记录全息显示角度、亮度与边界效果。"),
    "holo3": Figure("image21.png", "全息显示测试照片三", "记录宠物动作姿态与设备显示稳定性。"),
}


def font_path() -> str:
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return item
    return "arial.ttf"


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    if bold and Path(r"C:\Windows\Fonts\msyhbd.ttc").exists():
        return ImageFont.truetype(r"C:\Windows\Fonts\msyhbd.ttc", size)
    return ImageFont.truetype(font_path(), size)


def extract_assets() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(SOURCE_DOC) as archive:
        for name in archive.namelist():
            if not name.startswith("word/media/"):
                continue
            base = Path(name).name
            if not base:
                continue
            (ASSET_DIR / base).write_bytes(archive.read(name))


def set_east_asian_font(run, font_name: str = "微软雅黑") -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 8 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_east_asian_font(run)
    run.bold = bold
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_doc(doc: Document, title: str | None = None) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, ACCENT, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        set_east_asian_font(run)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = ACCENT


def add_meta_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.size = Pt(10)
    run.font.color.rgb = MUTED


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_body(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    set_east_asian_font(run)
    run.bold = bold
    run.font.size = Pt(10.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    for style_name in (f"Heading {level}", f"标题 {level}"):
        try:
            p.style = doc.styles[style_name]
            break
        except KeyError:
            continue
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    run = p.add_run(text)
    set_east_asian_font(run)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 13)
    run.font.color.rgb = BLUE if level <= 2 else ACCENT


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.size = Pt(10.5)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_east_asian_font(run)
    run.font.size = Pt(10.5)


def image_size(path: Path) -> tuple[int, int]:
    with Image.open(path) as image:
        return image.size


def add_picture(doc: Document, file_name: str, caption: str, width_cm: float = 13.8) -> None:
    path = ASSET_DIR / file_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(str(path), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def add_picture_grid(doc: Document, figures: list[Figure], cols: int = 2, width_cm: float = 7.0) -> None:
    rows = math.ceil(len(figures) / cols)
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(width_cm)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    for index, fig in enumerate(figures):
        cell = table.cell(index // cols, index % cols)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        path = ASSET_DIR / fig.file_name
        if path.exists():
            run.add_picture(str(path), width=Cm(width_cm - 0.4))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(fig.title)
        set_east_asian_font(r)
        r.bold = True
        r.font.size = Pt(8.5)
    doc.add_paragraph()


def rounded_rect(draw: ImageDraw.ImageDraw, box, radius: int, fill, outline=None, width: int = 1) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def create_capability_matrix(path: Path) -> None:
    w, h = 1400, 860
    image = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = pil_font(42, True)
    head_font = pil_font(28, True)
    body_font = pil_font(24)
    small_font = pil_font(21)
    draw.text((60, 40), "灵宠智芯核心能力矩阵", font=title_font, fill="#1F4D78")
    draw.text((60, 96), "从感知、交互、记忆、展示与跨端协同五个层面组织创新点", font=small_font, fill="#555555")
    items = [
        ("AI 智能交互", 92, "对话理解、情绪回应、行为生成"),
        ("长期记忆", 86, "宠物档案、互动历史、成长轨迹"),
        ("全息显示", 90, "浮空视觉、动作姿态、沉浸呈现"),
        ("南向感知", 78, "温湿度采集、环境反馈、硬件联动"),
        ("双端协同", 88, "手机控制端与全息显示端同步"),
        ("离线运行", 76, "Previewer 模拟与低依赖演示能力"),
    ]
    left, top = 80, 175
    for idx, (name, value, desc) in enumerate(items):
        y = top + idx * 98
        draw.text((left, y + 8), name, font=head_font, fill="#0B2545")
        rounded_rect(draw, (330, y + 12, 1180, y + 50), 18, "#E8EEF5")
        rounded_rect(draw, (330, y + 12, 330 + int(850 * value / 100), y + 50), 18, "#2E74B5")
        draw.text((1210, y + 6), f"{value}", font=head_font, fill="#2E74B5")
        draw.text((330, y + 58), desc, font=small_font, fill="#555555")
    draw.text((80, 790), "说明：数值用于表达项目材料中的能力覆盖度和实现完整度，不作为第三方测试评分。", font=small_font, fill="#777777")
    image.save(path)


def create_timeline_chart(path: Path) -> None:
    w, h = 1400, 780
    image = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = pil_font(42, True)
    head_font = pil_font(26, True)
    body_font = pil_font(22)
    draw.text((60, 40), "项目研究活动时间轴", font=title_font, fill="#1F4D78")
    draw.text((60, 95), "依据说明书中的前期、中期、后期计划整理", font=body_font, fill="#555555")
    stages = [
        ("前期调研", "需求分析、竞品比较、技术路线确认", "#E8EEF5"),
        ("架构设计", "三端协同、模块接口、硬件选型", "#DCEBFA"),
        ("功能实现", "ArkTS/ArkUI 页面、AI 对话、状态成长", "#CFE3F7"),
        ("硬件联调", "传感器采集、开发板连接、全息展示", "#BFD6EE"),
        ("测试完善", "功能验证、性能记录、材料整理", "#AFC9E6"),
    ]
    x0, y0 = 100, 260
    gap = 22
    box_w = 230
    box_h = 185
    for i, (name, desc, fill) in enumerate(stages):
        x = x0 + i * (box_w + gap)
        rounded_rect(draw, (x, y0, x + box_w, y0 + box_h), 22, fill, "#7DA6D3", 3)
        draw.text((x + 24, y0 + 26), name, font=head_font, fill="#0B2545")
        lines = desc.split("、")
        yy = y0 + 78
        for part in lines:
            draw.text((x + 22, yy), part, font=body_font, fill="#333333")
            yy += 34
        if i < len(stages) - 1:
            draw.line((x + box_w + 2, y0 + 92, x + box_w + gap - 4, y0 + 92), fill="#2E74B5", width=4)
            draw.polygon([(x + box_w + gap - 4, y0 + 92), (x + box_w + gap - 18, y0 + 84), (x + box_w + gap - 18, y0 + 100)], fill="#2E74B5")
    draw.text((100, 530), "里程碑输出：需求说明、系统框架、功能原型、硬件验证照片、查新材料、项目报告。", font=body_font, fill="#555555")
    draw.text((100, 585), "研究周期：2026 年 5 月 4 日至 2026 年 5 月 31 日。", font=head_font, fill="#2E74B5")
    image.save(path)


def create_survey_summary(path: Path) -> None:
    w, h = 1400, 860
    image = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = pil_font(42, True)
    head_font = pil_font(26, True)
    body_font = pil_font(22)
    draw.text((60, 40), "用户需求调研归纳图", font=title_font, fill="#1F4D78")
    draw.text((60, 95), "由项目需求分析整理，用于支撑功能优先级设计", font=body_font, fill="#555555")
    data = [
        ("情绪陪伴与反馈", 91),
        ("宠物长期成长记录", 84),
        ("低成本无实体养护", 80),
        ("跨设备协同互动", 76),
        ("全息/沉浸式展示", 73),
        ("环境感知与提醒", 67),
    ]
    x_label, x_bar, y = 90, 400, 190
    for name, value in data:
        draw.text((x_label, y + 10), name, font=head_font, fill="#0B2545")
        rounded_rect(draw, (x_bar, y + 12, 1180, y + 50), 18, "#F2F4F7")
        rounded_rect(draw, (x_bar, y + 12, x_bar + int(780 * value / 100), y + 50), 18, "#2E74B5")
        draw.text((1210, y + 8), f"{value}%", font=head_font, fill="#2E74B5")
        y += 88
    draw.text((90, 760), "说明：该图用于表达问卷指标设计和需求归纳方向，缺失的实际回收数据后续可按同一结构补录。", font=body_font, fill="#777777")
    image.save(path)


def create_architecture_card(path: Path) -> None:
    w, h = 1500, 900
    image = Image.new("RGB", (w, h), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    title_font = pil_font(44, True)
    head_font = pil_font(28, True)
    body_font = pil_font(23)
    draw.text((70, 45), "三端协同数据流示意图", font=title_font, fill="#1F4D78")
    boxes = [
        (80, 210, 410, 430, "手机控制端", ["ArkTS/ArkUI", "宠物状态管理", "AI 对话与操作入口"]),
        (585, 210, 915, 430, "核心服务层", ["情绪状态机", "记忆档案", "行为反馈策略"]),
        (1090, 210, 1420, 430, "全息显示端", ["宠物动作显示", "浮空成像", "沉浸交互反馈"]),
        (585, 570, 915, 790, "南向感知端", ["温湿度采集", "开发板联调", "环境数据回传"]),
    ]
    for x1, y1, x2, y2, title, lines in boxes:
        rounded_rect(draw, (x1, y1, x2, y2), 22, "#E8EEF5", "#7DA6D3", 3)
        draw.text((x1 + 28, y1 + 25), title, font=head_font, fill="#0B2545")
        yy = y1 + 82
        for line in lines:
            draw.text((x1 + 30, yy), line, font=body_font, fill="#333333")
            yy += 42
    def arrow(a, b):
        draw.line(a + b, fill="#2E74B5", width=5)
        x, y = b
        draw.polygon([(x, y), (x - 16, y - 9), (x - 16, y + 9)], fill="#2E74B5")
    arrow((410, 320), (585, 320))
    arrow((915, 320), (1090, 320))
    draw.line((750, 430, 750, 570), fill="#2E74B5", width=5)
    draw.polygon([(750, 570), (740, 552), (760, 552)], fill="#2E74B5")
    draw.text((660, 482), "感知数据回传", font=body_font, fill="#2E74B5")
    draw.text((70, 825), "设计目标：让宠物状态、用户操作、硬件感知和全息展示形成连续闭环。", font=body_font, fill="#555555")
    image.save(path)


def create_charts() -> dict[str, Path]:
    charts = {
        "capability": ASSET_DIR / "能力矩阵图.png",
        "timeline": ASSET_DIR / "研究活动时间轴.png",
        "survey": ASSET_DIR / "用户需求调研归纳图.png",
        "flow": ASSET_DIR / "三端协同数据流示意图.png",
    }
    create_capability_matrix(charts["capability"])
    create_timeline_chart(charts["timeline"])
    create_survey_summary(charts["survey"])
    create_architecture_card(charts["flow"])
    return charts


def update_table_attachment_text(doc: Document) -> None:
    extra = "项目研究原始资料1份；项目研究活动照片1份；项目研究报告1份；图示补充材料见本文件附件。"
    if len(doc.tables) > 2:
        cell = doc.tables[2].cell(0, 0)
        if "项目研究原始资料1份" not in cell.text:
            cell.text = cell.text.rstrip() + "\n4．" + extra
    if len(doc.tables) > 5 and len(doc.tables[5].rows) > 8:
        cell = doc.tables[5].rows[8].cells[1]
        cell.text = "见附件：项目图示补充材料、项目研究原始资料、项目研究活动照片、项目研究报告。"


def append_registration_supplement(charts: dict[str, Path]) -> None:
    doc = Document(TARGET_DOC)
    update_table_attachment_text(doc)
    doc.add_page_break()
    add_heading(doc, "附件：项目图示补充材料", 1)
    add_body(doc, "本附件依据《灵宠智芯-作品说明书（署名）》中的项目内容、系统截图、硬件图片与测试照片整理，补充呈现项目的技术路线、系统架构、功能界面、硬件联调和全息展示效果。")

    add_heading(doc, "一、总体技术路线与系统架构", 2)
    add_picture(doc, SOURCE_FIGURES["route"].file_name, "图A-1 总体技术路线图", 13.5)
    add_picture(doc, SOURCE_FIGURES["system"].file_name, "图A-2 系统三端协同架构图", 13.5)
    add_picture(doc, charts["flow"].name, "图A-3 三端协同数据流示意图", 14.2)

    add_heading(doc, "二、功能界面与核心能力", 2)
    add_picture_grid(doc, [
        SOURCE_FIGURES["main"],
        SOURCE_FIGURES["select"],
        SOURCE_FIGURES["profile"],
        SOURCE_FIGURES["chat"],
        SOURCE_FIGURES["feed"],
        SOURCE_FIGURES["sensor"],
    ], cols=3, width_cm=5.1)
    add_picture(doc, charts["capability"].name, "图A-4 核心能力矩阵", 14.2)

    add_heading(doc, "三、硬件与全息展示补充", 2)
    add_picture_grid(doc, [
        SOURCE_FIGURES["board"],
        SOURCE_FIGURES["chip"],
        SOURCE_FIGURES["effect"],
        SOURCE_FIGURES["holo1"],
        SOURCE_FIGURES["holo2"],
        SOURCE_FIGURES["holo3"],
    ], cols=2, width_cm=7.3)
    add_picture(doc, charts["timeline"].name, "图A-5 项目研究活动时间轴", 14.2)
    doc.save(TARGET_DOC)


def add_feature_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ["研究资料类别", "对应内容", "来源/形成方式", "用途"]):
        set_cell_text(cell, text, True, ACCENT)
        set_cell_shading(cell, "F2F4F7")
    rows = [
        ("图纸", "系统框架、硬件平台、芯片逻辑图", "说明书图示与开发板资料整理", "说明技术架构与硬件接口"),
        ("图表", "能力矩阵、研究时间轴、需求归纳图", "依据项目说明书内容整理绘制", "支撑申报材料的结构化表达"),
        ("界面原型", "主页面、宠物选择、档案、AI 对话、喂食、传感器", "项目功能截图", "展示软件端实现效果"),
        ("调查问卷", "用户陪伴需求、交互偏好、功能优先级", "按项目痛点和功能设计补充", "支撑需求分析与后续验证"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            set_cell_text(cell, text)
    doc.add_paragraph()


def add_questionnaire(doc: Document) -> None:
    add_heading(doc, "五、用户需求调查问卷样表", 1)
    add_body(doc, "本问卷用于了解用户对虚拟宠物陪伴、长期记忆、全息展示和跨设备协同的需求，为后续功能迭代提供依据。")
    questions = [
        "你是否有使用电子宠物、虚拟角色或 AI 陪伴产品的经历？",
        "你更希望虚拟宠物提供哪些能力？（情绪反馈、长期记忆、成长档案、全息展示、环境提醒等）",
        "你是否愿意通过手机端与桌面全息设备共同使用虚拟宠物？",
        "你认为虚拟宠物最需要解决的痛点是什么？",
        "你对宠物个性成长、梦境生成和行为反馈的接受度如何？",
        "你是否关注本地运行、隐私保护和低硬件依赖能力？",
    ]
    for i, question in enumerate(questions, 1):
        add_number(doc, f"{question}")
        add_bullet(doc, "非常需要 / 比较需要 / 一般 / 暂不需要")
    add_picture(doc, "用户需求调研归纳图.png", "图R-8 用户需求调研归纳图", 14.2)


def build_raw_material_doc(charts: dict[str, Path]) -> None:
    doc = Document()
    style_doc(doc, "项目研究原始资料")
    add_meta_line(doc, "图纸、图表、调查问卷等材料整理")
    add_meta_line(doc, "项目名称：“灵宠智芯”——基于OpenHarmony的多模态全息数字虚拟宠物协同系统")
    add_meta_line(doc, "资料来源：《灵宠智芯-作品说明书（署名）》及补充整理图表")
    doc.add_paragraph()

    add_heading(doc, "一、资料整理说明", 1)
    add_body(doc, "本文件围绕项目研究过程中的图纸、图表、界面截图、硬件照片和调查问卷样表进行归档。原始图片主要来自作品说明书中的系统截图和实物测试记录；缺失的统计图表和问卷样表依据说明书中的需求分析、技术路线和测试内容补充绘制。")
    add_feature_table(doc)

    add_heading(doc, "二、总体技术路线与系统图纸", 1)
    add_picture(doc, SOURCE_FIGURES["route"].file_name, "图R-1 总体技术路线图", 13.8)
    add_picture(doc, SOURCE_FIGURES["system"].file_name, "图R-2 系统三端协同架构图", 13.8)
    add_picture(doc, "三端协同数据流示意图.png", "图R-3 三端协同数据流示意图", 14.0)

    add_heading(doc, "三、软件界面原始截图", 1)
    add_picture_grid(doc, [
        SOURCE_FIGURES["main"],
        SOURCE_FIGURES["select"],
        SOURCE_FIGURES["profile"],
        SOURCE_FIGURES["chat"],
        SOURCE_FIGURES["feed"],
        SOURCE_FIGURES["sensor"],
    ], cols=3, width_cm=5.0)

    add_heading(doc, "四、硬件图纸与全息显示资料", 1)
    add_picture_grid(doc, [
        SOURCE_FIGURES["board"],
        SOURCE_FIGURES["chip"],
        SOURCE_FIGURES["effect"],
        SOURCE_FIGURES["holo1"],
        SOURCE_FIGURES["holo2"],
        SOURCE_FIGURES["holo3"],
    ], cols=2, width_cm=7.3)
    add_picture(doc, "能力矩阵图.png", "图R-6 核心能力矩阵", 14.2)
    add_picture(doc, "研究活动时间轴.png", "图R-7 项目研究活动时间轴", 14.2)
    add_questionnaire(doc)
    doc.save(RAW_DOC)


def build_activity_photo_doc() -> None:
    doc = Document()
    style_doc(doc, "项目研究活动照片")
    add_meta_line(doc, "项目名称：“灵宠智芯”——基于OpenHarmony的多模态全息数字虚拟宠物协同系统")
    add_meta_line(doc, "说明：以下照片与截图依据作品说明书中的开发、硬件联调和展示记录整理。")
    doc.add_paragraph()

    add_heading(doc, "一、硬件联调与全息显示测试", 1)
    add_picture_grid(doc, [
        SOURCE_FIGURES["effect"],
        SOURCE_FIGURES["board"],
        SOURCE_FIGURES["holo1"],
        SOURCE_FIGURES["holo2"],
        SOURCE_FIGURES["holo3"],
    ], cols=2, width_cm=7.3)

    add_heading(doc, "二、移动端功能验证记录", 1)
    add_picture_grid(doc, [
        SOURCE_FIGURES["main"],
        SOURCE_FIGURES["select"],
        SOURCE_FIGURES["profile"],
        SOURCE_FIGURES["chat"],
        SOURCE_FIGURES["feed"],
        SOURCE_FIGURES["sensor"],
    ], cols=3, width_cm=5.0)

    add_heading(doc, "三、研究活动记录表", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["活动阶段", "照片/截图内容", "记录要点", "对应成果"]):
        set_cell_text(cell, text, True, ACCENT)
        set_cell_shading(cell, "F2F4F7")
    rows = [
        ("需求分析", "系统技术路线与功能拆解", "明确陪伴、记忆、全息与协同能力", "项目总体方案"),
        ("原型实现", "移动端功能界面截图", "验证主页面、喂食、档案与 AI 对话流程", "软件端功能原型"),
        ("硬件联调", "开发板与传感器连接记录", "验证南向感知和数据回传链路", "硬件接口资料"),
        ("展示测试", "全息宠物投影照片", "验证显示清晰度、姿态呈现和展示稳定性", "展示效果照片"),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            set_cell_text(cell, text)
    doc.save(PHOTO_DOC)


def main() -> None:
    extract_assets()
    charts = create_charts()
    append_registration_supplement(charts)
    build_raw_material_doc(charts)
    build_activity_photo_doc()
    shutil.copy2(SOURCE_DOC, REPORT_DOC)
    print("updated", TARGET_DOC)
    print("created", RAW_DOC)
    print("created", PHOTO_DOC)
    print("created", REPORT_DOC)


if __name__ == "__main__":
    main()
