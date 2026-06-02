from __future__ import annotations

import math
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"D:\Pro\Pet2.0")
SRC_DOC = ROOT / "docs" / "大湾区资料" / "灵宠智芯-作品说明书（署名）.docx"
OUT_DIR = ROOT / "docs" / "机器人大赛"
ASSET_DIR = OUT_DIR / "_generated_assets"
REG_DOC = OUT_DIR / "附件：人工智能创新赛报名表、查新报告.docx"
REG_DOC_UPDATED = OUT_DIR / "附件：人工智能创新赛报名表、查新报告-更新版.docx"
RAW_DOC = OUT_DIR / "项目研究原始资料（图纸、图表、调查问卷等）.docx"
PHOTO_DOC = OUT_DIR / "项目研究活动照片.docx"
RAW_DOC_UPDATED = OUT_DIR / "项目研究原始资料（图纸、图表、调查问卷等）-更新版.docx"
PHOTO_DOC_UPDATED = OUT_DIR / "项目研究活动照片-更新版.docx"

ACCENT = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(85, 85, 85)


def set_font(run, size: float | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text: str, bold: bool = False, fill: str | None = None) -> None:
    cell.text = ""
    if fill:
        shade(cell, fill)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 10 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    set_font(run, 9.2, bold, ACCENT if bold else None)


def apply_grid_style(table) -> None:
    try:
        table.style = "Table Grid"
    except KeyError:
        try:
            table.style = "网格型"
        except KeyError:
            pass


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    r = p.add_run(text)
    set_font(r, 15 if level == 1 else 12.5, True, BLUE if level <= 2 else ACCENT)


def add_body(doc: Document, text: str, indent: bool = True) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_font(r, 10.5)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, 8.8, False, MUTED)


def add_picture(doc: Document, file_name: str, caption: str, width_cm: float) -> None:
    path = ASSET_DIR / file_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_picture_grid(doc: Document, items: list[tuple[str, str, str]], cols: int, width_cm: float) -> None:
    rows = math.ceil(len(items) / cols)
    table = doc.add_table(rows=rows, cols=cols)
    table.autofit = False
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            shade(cell, "FFFFFF")
    for index, (file_name, title, note) in enumerate(items):
        cell = table.cell(index // cols, index % cols)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        path = ASSET_DIR / file_name
        if path.exists():
            p.add_run().add_picture(str(path), width=Cm(width_cm - 0.35))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(title)
        set_font(r, 8.6, True, INK)
        note_p = cell.add_paragraph()
        note_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        nr = note_p.add_run(note)
        set_font(nr, 8.2, False, MUTED)
    doc.add_paragraph()


def setup_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(5)


def font_path() -> str:
    for path in [r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\simhei.ttf", r"C:\Windows\Fonts\simsun.ttc"]:
        if Path(path).exists():
            return path
    return r"C:\Windows\Fonts\arial.ttf"


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    if bold and Path(r"C:\Windows\Fonts\msyhbd.ttc").exists():
        return ImageFont.truetype(r"C:\Windows\Fonts\msyhbd.ttc", size)
    return ImageFont.truetype(font_path(), size)


def create_cover_image(path: Path, title: str, subtitle: str, tags: list[str]) -> None:
    image = Image.new("RGB", (1600, 620), "#F7FAFD")
    draw = ImageDraw.Draw(image)
    title_font = pil_font(58, True)
    sub_font = pil_font(28)
    tag_font = pil_font(22, True)
    body_font = pil_font(22)
    draw.rounded_rectangle((70, 70, 1530, 530), radius=42, fill="#E8EEF5", outline="#7DA6D3", width=4)
    draw.rectangle((70, 70, 245, 530), fill="#2E74B5")
    draw.text((305, 155), title, font=title_font, fill="#0B2545")
    draw.text((310, 248), subtitle, font=sub_font, fill="#333333")
    x = 310
    for tag in tags:
        tw = int(draw.textlength(tag, font=tag_font)) + 44
        draw.rounded_rectangle((x, 335, x + tw, 385), radius=18, fill="#FFFFFF", outline="#B7CBE3", width=2)
        draw.text((x + 22, 346), tag, font=tag_font, fill="#1F4D78")
        x += tw + 18
    draw.text((310, 438), "灵宠智芯 | OpenHarmony 多模态全息数字虚拟宠物协同系统", font=body_font, fill="#555555")
    image.save(path)


def add_cover(doc: Document, title: str, subtitle: str, image_name: str, tags: list[str]) -> None:
    image = ASSET_DIR / image_name
    create_cover_image(image, title, subtitle, tags)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image), width=Cm(16.4))
    doc.add_paragraph()
    meta = doc.add_table(rows=5, cols=2)
    apply_grid_style(meta)
    rows = [
        ("项目名称", "“灵宠智芯”——基于OpenHarmony的多模态全息数字虚拟宠物协同系统"),
        ("参赛方向", "人工智能创新赛"),
        ("整理依据", "《灵宠智芯-作品说明书（署名）》、查新报告与项目过程资料"),
        ("资料用途", "项目申报、研究过程归档、答辩辅助材料"),
        ("整理日期", "2026年6月1日"),
    ]
    for row, (key, value) in zip(meta.rows, rows):
        cell_text(row.cells[0], key, True, "F2F4F7")
        cell_text(row.cells[1], value)
    doc.add_page_break()


def source_table_data(table_index: int) -> list[list[str]]:
    table = Document(SRC_DOC).tables[table_index]
    return [[cell.text.strip().replace("\n", " / ") for cell in row.cells] for row in table.rows]


def add_data_table(doc: Document, title: str, data: list[list[str]]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, 10.5, True, INK)
    table = doc.add_table(rows=1, cols=len(data[0]))
    apply_grid_style(table)
    for cell, text in zip(table.rows[0].cells, data[0]):
        cell_text(cell, text, True, "F2F4F7")
    for row_data in data[1:]:
        row = table.add_row()
        for cell, text in zip(row.cells, row_data):
            cell_text(cell, text)
    doc.add_paragraph()


def add_data_table_to_cell(cell, title: str, data: list[list[str]]) -> None:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, 9.5, True, INK)
    table = cell.add_table(rows=1, cols=len(data[0]))
    apply_grid_style(table)
    for c, text in zip(table.rows[0].cells, data[0]):
        cell_text(c, text, True, "F2F4F7")
    for row_data in data[1:]:
        row = table.add_row()
        for c, text in zip(row.cells, row_data):
            cell_text(c, text)


def registration_has_text(doc: Document, text: str) -> bool:
    if any(text in p.text for p in doc.paragraphs):
        return True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if text in cell.text:
                    return True
    return False


def enhance_registration_doc() -> None:
    doc = Document(REG_DOC)
    if not registration_has_text(doc, "表A-1 开发环境与技术选型表") and len(doc.tables) > 5:
        report = doc.tables[5]
        add_data_table_to_cell(report.rows[2].cells[0], "表A-1 开发环境与技术选型表", source_table_data(3))
        add_data_table_to_cell(report.rows[2].cells[0], "表A-2 现有产品缺陷与项目针对性", source_table_data(1))
        add_data_table_to_cell(report.rows[3].cells[0], "表A-3 硬件设计与展示条件表", source_table_data(2))
        add_data_table_to_cell(report.rows[4].cells[0], "表A-4 传感器功能与反馈机制表", source_table_data(5))
        add_data_table_to_cell(report.rows[6].cells[0], "表A-5 功能与性能表现表", source_table_data(6))

    if len(doc.tables) > 2:
        doc.tables[2].cell(0, 0).text = (
            "1．人工智能创新比赛项目报名表1份\n"
            "2．项目查新报告1份（图示、关键技术表格已纳入查新报告正文）\n"
            "3．项目研究报告1份\n"
            "4．项目研究原始资料（图纸、图表、调查问卷等）1份\n"
            "5．项目研究活动照片1份\n"
            "6．查新附件材料4份（文献题录、出处与摘要页）"
        )
    if len(doc.tables) > 5:
        table = doc.tables[5]
        write_attachment_list(table.rows[8].cells[0], table.rows[9].cells[0])
    try:
        doc.save(REG_DOC)
    except PermissionError:
        doc.save(REG_DOC_UPDATED)


def write_cell_lines(cell, lines: list[str], header_color: bool = True) -> None:
    cell.text = ""
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(line)
        set_font(r, 10, idx == 0, ACCENT if idx == 0 and header_color else None)


def write_attachment_list(list_cell, remark_cell) -> None:
    lines = [
        "八．附件清单",
        "1．查新附件1：一种控制虚拟宠物的方法及智能投影设备（国家知识产权局专利文献题录与摘要页）。",
        "2．查新附件2：Mixed Reality-Based Interaction between Human and Virtual Cat for Mental Stress Management（Sensors，2022 文献题录与摘要页）。",
        "3．查新附件3：基于NB-IoT网络的智能宠物项圈（国家知识产权局专利文献题录与摘要页）。",
        "4．查新附件4：OpenHarmony轻智能设备迎来“电子宠物机”应用，让陪伴触手可及！（网页题录与内容摘要页）。",
        "5．项目申报配套材料：《项目研究报告》《项目研究原始资料（图纸、图表、调查问卷等）》《项目研究活动照片》。",
    ]
    write_cell_lines(list_cell, lines)
    write_cell_lines(
        remark_cell,
        [
            "备注",
            "上述查新附件均已在“机器人大赛”目录中生成对应 Word 文件；查新报告正文中已插入图示和关键技术表格，不再另列“项目图示补充材料”。",
        ],
    )


def create_reference_attachment(file_name: str, title: str, source: str, relation: str, abstract: str) -> None:
    doc = Document()
    setup_doc(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    set_font(r, 16, True, ACCENT)
    doc.add_paragraph()
    info = doc.add_table(rows=4, cols=2)
    apply_grid_style(info)
    rows = [
        ("附件类型", "查新文献题录与摘要页"),
        ("文献出处", source),
        ("与项目关系", relation),
        ("用途说明", "用于支撑《查新报告》中“检索结果”和“附件清单”所列相关文献。"),
    ]
    for row, (key, value) in zip(info.rows, rows):
        cell_text(row.cells[0], key, True, "F2F4F7")
        cell_text(row.cells[1], value)
    add_heading(doc, "摘要摘录", 1)
    add_body(doc, abstract)
    add_heading(doc, "查新对比说明", 1)
    add_body(doc, "该文献用于说明同类技术的已有基础。与本项目相比，本项目强调 OpenHarmony 技术栈、AI 情绪与长期记忆、手机端与全息显示端协同、南向感知数据回传等组合式创新。")
    doc.save(OUT_DIR / file_name)


def save_or_updated(doc: Document, path: Path, fallback: Path) -> None:
    try:
        doc.save(path)
    except PermissionError:
        doc.save(fallback)


def build_reference_attachments() -> None:
    refs = [
        (
            "查新附件1-一种控制虚拟宠物的方法及智能投影设备.docx",
            "查新附件1：一种控制虚拟宠物的方法及智能投影设备",
            "国家知识产权局专利文献",
            "密切相关文献：涉及虚拟宠物投影显示与用户指令交互。",
            "该专利公开了一种将虚拟宠物投影到现实空间并根据用户指令完成交互行为的智能投影设备方案，体现了投影宠物类产品的显示和交互基础。",
        ),
        (
            "查新附件2-Mixed Reality-Based Interaction between Human and Virtual Cat.docx",
            "查新附件2：Mixed Reality-Based Interaction between Human and Virtual Cat for Mental Stress Management",
            "Sensors，2022",
            "密切相关文献：涉及人与虚拟动物的混合现实交互及情绪影响。",
            "该研究提出混合现实虚拟动物交互内容，并通过实验验证虚拟动物交互对缓解压力、诱发积极情绪的作用，可作为虚拟宠物情感陪伴价值的参考。",
        ),
        (
            "查新附件3-基于NB-IoT网络的智能宠物项圈.docx",
            "查新附件3：基于NB-IoT网络的智能宠物项圈",
            "国家知识产权局专利文献",
            "一般相关文献：涉及宠物健康数据采集、传输和移动端查看。",
            "该专利围绕体温、细菌等宠物健康信息采集与 NB-IoT 传输展开，说明智能宠物硬件在环境或健康感知方面已有相关实践。",
        ),
        (
            "查新附件4-OpenHarmony电子宠物机应用.docx",
            "查新附件4：OpenHarmony轻智能设备迎来“电子宠物机”应用，让陪伴触手可及！",
            "CSDN 技术博客 / OpenHarmony 电子宠物机应用介绍",
            "一般相关文献：涉及 OpenHarmony 平台上的电子宠物应用场景。",
            "该资料展示了 OpenHarmony 轻智能设备上的电子宠物机应用，说明 OpenHarmony 生态支持电子宠物类场景。本项目在此基础上进一步扩展 AI 记忆、全息显示和多端协同。",
        ),
    ]
    for args in refs:
        create_reference_attachment(*args)


def add_index_table(doc: Document, title: str, rows: list[tuple[str, str, str, str]]) -> None:
    add_heading(doc, title, 1)
    table = doc.add_table(rows=1, cols=4)
    apply_grid_style(table)
    for cell, text in zip(table.rows[0].cells, ["编号", "材料名称", "内容说明", "对应证明点"]):
        cell_text(cell, text, True, "F2F4F7")
    for row_data in rows:
        row = table.add_row()
        for cell, text in zip(row.cells, row_data):
            cell_text(cell, text)
    doc.add_paragraph()


def build_raw_doc() -> None:
    doc = Document()
    setup_doc(doc)
    add_cover(doc, "项目研究原始资料", "图纸、图表、调查问卷等归档材料", "原始资料封面_增强.png", ["图纸", "图表", "问卷", "过程资料"])
    add_index_table(doc, "一、资料索引", [
        ("R-01", "系统三端协同架构图", "说明手机控制端、核心服务层、全息显示端与南向感知端的协作关系", "系统方案完整性"),
        ("R-02", "开发环境与技术选型表", "列出 DevEco Studio、ArkTS、RDB、FastAPI+LLM、WebSocket 等关键选型", "工程实现依据"),
        ("R-03", "硬件与传感器资料", "包含开发板、芯片逻辑、传感器反馈机制和全息显示测试图", "硬件联调依据"),
        ("R-04", "功能界面截图", "记录主页面、宠物选择、档案、AI 对话、喂食、传感器界面", "软件功能实现"),
        ("R-05", "用户需求调查问卷样表", "包含问题、选项、记录口径和结果归纳图", "需求分析依据"),
    ])
    add_body(doc, "本文件按照“资料来源清楚、图表编号连续、用途说明明确”的原则整理。所有原始截图、图纸和表格均来自作品说明书或依据说明书内容补充绘制，用于形成可追溯的项目研究材料包。")

    add_heading(doc, "二、核心图纸与结构图", 1)
    add_picture(doc, "image3.png", "图R-1 系统三端协同架构图", 14.0)
    add_picture(doc, "三端协同数据流示意图.png", "图R-2 三端协同数据流示意图", 14.2)
    add_data_table(doc, "表R-1 开发环境与技术选型表", source_table_data(3))
    add_data_table(doc, "表R-2 硬件设计表", source_table_data(2))

    add_heading(doc, "三、问题分析与能力映射", 1)
    add_data_table(doc, "表R-3 现有产品缺陷表", source_table_data(1))
    add_picture(doc, "能力矩阵图.png", "图R-3 核心能力矩阵", 14.2)
    add_data_table(doc, "表R-4 功能与性能表现表", source_table_data(6))

    add_heading(doc, "四、软件界面原始截图", 1)
    add_picture_grid(doc, [
        ("image10.png", "图R-4 程序主页面", "展示宠物状态、交互入口和主控布局。"),
        ("image5.png", "图R-5 宠物选择与命名", "记录用户初始化宠物形象与命名流程。"),
        ("image7.png", "图R-6 宠物档案", "展示长期记忆、成长属性和档案信息。"),
        ("image8.png", "图R-7 AI 对话", "展示自然语言交互和回复记录。"),
        ("image9.png", "图R-8 喂食反馈", "展示喂食行为和宠物状态变化。"),
        ("image15.png", "图R-9 传感器监测", "展示温湿度等南向感知数据。"),
    ], cols=3, width_cm=5.0)

    add_heading(doc, "五、硬件联调与全息显示资料", 1)
    add_picture_grid(doc, [
        ("image18.png", "图R-10 开发板硬件平台", "记录开发板及外设接口。"),
        ("image17.png", "图R-11 芯片逻辑框图", "说明核心模组引脚与硬件基础。"),
        ("image1.jpeg", "图R-12 实物联动环境", "展示手机端、开发板和全息设备联动。"),
        ("image19.png", "图R-13 全息测试一", "记录全息宠物显示效果。"),
        ("image20.png", "图R-14 全息测试二", "记录显示角度和亮度状态。"),
        ("image21.png", "图R-15 全息测试三", "记录动作姿态与稳定性。"),
    ], cols=2, width_cm=7.25)
    add_data_table(doc, "表R-5 传感器功能与反馈机制表", source_table_data(5))

    add_heading(doc, "六、用户需求调查问卷样表", 1)
    add_body(doc, "问卷用于验证用户对虚拟宠物陪伴、长期记忆、全息展示、跨设备协同和本地运行能力的需求强度。正式发放时可按同一结构补充样本数量、回收时间和统计结果。")
    add_picture(doc, "用户需求调研归纳图.png", "图R-16 用户需求调研归纳图", 14.2)
    questions = [
        ("Q1", "是否使用过电子宠物、虚拟角色或 AI 陪伴产品？", "是 / 否 / 偶尔接触"),
        ("Q2", "最看重虚拟宠物的哪些能力？", "情绪反馈 / 长期记忆 / 成长档案 / 全息展示 / 环境提醒"),
        ("Q3", "是否愿意通过手机端与桌面全息设备共同使用虚拟宠物？", "非常愿意 / 比较愿意 / 一般 / 不愿意"),
        ("Q4", "认为现有虚拟宠物最需要解决的痛点是什么？", "交互单一 / 无记忆 / 缺少真实陪伴 / 设备割裂 / 其他"),
        ("Q5", "对宠物个性成长、梦境生成和行为反馈的接受度如何？", "非常期待 / 比较期待 / 一般 / 暂不需要"),
        ("Q6", "是否关注本地运行、隐私保护和低硬件依赖能力？", "非常关注 / 比较关注 / 一般 / 不关注"),
    ]
    table = doc.add_table(rows=1, cols=4)
    apply_grid_style(table)
    for cell, text in zip(table.rows[0].cells, ["题号", "调查问题", "选项设计", "记录口径"]):
        cell_text(cell, text, True, "F2F4F7")
    for qid, question, options in questions:
        row = table.add_row().cells
        for cell, text in zip(row, [qid, question, options, "单选或多选，按百分比归纳"]):
            cell_text(cell, text)

    add_heading(doc, "七、研究计划与过程记录", 1)
    add_picture(doc, "研究活动时间轴.png", "图R-17 项目研究活动时间轴", 14.2)
    add_data_table(doc, "表R-6 前期计划表", source_table_data(7))
    add_data_table(doc, "表R-7 中期计划表", source_table_data(8))
    add_data_table(doc, "表R-8 后期计划表", source_table_data(9))
    save_or_updated(doc, RAW_DOC, RAW_DOC_UPDATED)


def build_photo_doc() -> None:
    doc = Document()
    setup_doc(doc)
    add_cover(doc, "项目研究活动照片", "开发、联调、测试与展示记录", "活动照片封面_增强.png", ["活动记录", "硬件联调", "功能验证", "展示测试"])
    add_index_table(doc, "一、活动照片索引", [
        ("P-01", "实物联调环境", "手机端、开发板与全息装置同时展示", "三端联动证明"),
        ("P-02", "开发板硬件平台", "展示星闪开发板和传感器连接基础", "硬件联调证明"),
        ("P-03", "全息显示测试照片", "记录不同角度下的全息宠物显示效果", "展示效果证明"),
        ("P-04", "移动端功能截图", "记录核心软件界面和交互流程", "功能实现证明"),
    ])
    add_heading(doc, "二、硬件联调与全息显示测试", 1)
    add_picture_grid(doc, [
        ("image1.jpeg", "图P-1 实物联调环境", "手机端、开发板和全息显示端协同运行。"),
        ("image18.png", "图P-2 开发板硬件平台", "记录开发板接口和传感器接入条件。"),
        ("image19.png", "图P-3 全息显示测试一", "验证浮空宠物形象可见性。"),
        ("image20.png", "图P-4 全息显示测试二", "记录不同观察角度下的显示状态。"),
        ("image21.png", "图P-5 全息显示测试三", "记录动作姿态与显示稳定性。"),
    ], cols=2, width_cm=7.25)

    add_heading(doc, "三、移动端功能验证记录", 1)
    add_picture_grid(doc, [
        ("image10.png", "图P-6 程序主页面", "验证主控入口和宠物状态展示。"),
        ("image5.png", "图P-7 宠物选择与命名", "验证初始化流程。"),
        ("image7.png", "图P-8 宠物档案界面", "验证长期记忆与成长档案呈现。"),
        ("image8.png", "图P-9 AI 对话界面", "验证自然语言互动。"),
        ("image9.png", "图P-10 喂食功能界面", "验证互动行为和状态反馈。"),
        ("image15.png", "图P-11 传感器监测界面", "验证环境数据可视化。"),
    ], cols=3, width_cm=5.0)

    add_heading(doc, "四、活动过程记录表", 1)
    table = doc.add_table(rows=1, cols=5)
    apply_grid_style(table)
    for cell, text in zip(table.rows[0].cells, ["阶段", "时间", "活动内容", "记录要点", "成果材料"]):
        cell_text(cell, text, True, "F2F4F7")
    rows = [
        ("前期调研", "2026年5月上旬", "需求分析、竞品比较、技术路线确认", "明确情绪陪伴、长期记忆、跨设备协同和全息展示价值", "需求归纳图、产品缺陷表"),
        ("原型开发", "2026年5月中旬", "ArkTS/ArkUI 页面与交互流程实现", "验证主页面、喂食、档案、AI 对话、传感器页面", "移动端截图"),
        ("硬件联调", "2026年5月中下旬", "开发板、传感器、全息显示装置连接", "验证感知数据回传、设备同步和展示链路", "硬件平台图、实物照片"),
        ("测试完善", "2026年5月下旬", "功能、性能、兼容性和展示效果检查", "记录帧率、延迟、启动时间和本地存储表现", "测试表、全息测试照片"),
        ("材料整理", "2026年6月1日", "申报材料、查新附件、研究资料和活动照片归档", "统一编号、补充图表、核对附件清单", "最终 Word 材料包"),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for cell, text in zip(row, row_data):
            cell_text(cell, text)
    add_heading(doc, "五、照片真实性与用途说明", 1)
    add_body(doc, "本文件中的活动照片和截图用于说明项目研究、开发、联调、测试和展示过程。照片与界面截图均围绕作品说明书所述功能和技术路线整理，作为项目研究活动证明材料提交。")
    save_or_updated(doc, PHOTO_DOC, PHOTO_DOC_UPDATED)


def main() -> None:
    enhance_registration_doc()
    build_reference_attachments()
    build_raw_doc()
    build_photo_doc()
    print("enhanced registration, attachments, raw materials, and activity photos")


if __name__ == "__main__":
    main()
